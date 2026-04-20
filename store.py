from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

# Central path config — change this one line to move the output folder
OUTPUT_DIR = r"C:\ub"


class Store:
    def __init__(self):
        pass

    def docc(self, date):
        self.doc = Document()
        self.data = []
        self.data1 = []
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        self.doc.add_heading(f"Sales and Expenditure report on {date}", level=1)

    def opening_balance(self, date):
        path = os.path.join(OUTPUT_DIR, f"{date}.docx")
        try:
            doc1 = Document(path)
            table = doc1.tables[3]
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            strPrevDayTotal = rows[3][1]
            self.prevDayTotal = strPrevDayTotal.partition("=")[2].strip()
        except Exception as e:
            print(f"[opening_balance] Could not read previous file: {e}")
            self.prevDayTotal = "0"

        table0 = self.doc.add_table(rows=1, cols=2)
        table0.style = "Table Grid"
        table0.cell(0, 0).text = "Opening Balance"
        table0.cell(0, 1).text = self.prevDayTotal

    def sales(self, cash, upi, credit, profit):
        self.profit = profit
        self.doc.add_heading("Sales", level=1)
        self.data1 = [
            ["Cash Sales", cash],
            ["Phone Pay Sales", upi],
            ["Credit sales", credit],
            [" ", " "],
            ["Total Sales", self.profit],
        ]
        table1 = self.doc.add_table(rows=len(self.data1), cols=2)
        table1.style = "Table Grid"
        for i, row in enumerate(self.data1):
            for j, cell in enumerate(row):
                table1.cell(i, j).text = str(cell)

    def expense(self, data2, onlineLoss, cashLoss, loss):
        """Note: data2 is already a copy — safe to mutate here."""
        self.loss = loss
        self.doc.add_heading("Purchase and Expense", level=1)
        data2.append(["", "", "", ""])
        data2.append(["Online Expenses", onlineLoss, "Cash Expenses", cashLoss])
        data2.append(["", "", "", ""])
        data2.append(["Total Expense", "", "", self.loss])
        table2 = self.doc.add_table(rows=len(data2), cols=4)
        table2.style = "Table Grid"
        for i, row in enumerate(data2):
            for j, cell in enumerate(row):
                table2.cell(i, j).text = str(cell)

    def net_sales(self, account, hand, total):
        self.doc.add_heading("Net Balance", level=1)
        table3 = self.doc.add_table(rows=4, cols=2)
        strDayTotal = f"{self.profit} - {self.loss} = {total}"
        grandTotal = int(self.prevDayTotal) + total
        strGrandTotal = f"{total} + {self.prevDayTotal} = {grandTotal}"
        table3.style = "Table Grid"
        table3.cell(0, 0).text = "Per day Balance"
        table3.cell(0, 1).text = strDayTotal
        table3.cell(1, 0).text = "Cash in Hand"
        table3.cell(1, 1).text = hand
        table3.cell(2, 0).text = "Cash at Bank"
        table3.cell(2, 1).text = account
        table3.cell(3, 0).text = "Closing Balance"
        table3.cell(3, 1).text = strGrandTotal

    def info(self, date, net_day_profit):
        """Update the running summary tracker (data.docx).

        data.docx structure:
          Heading: "Summary as on <date>"
          Table (3 rows x 2 cols):
            Row 0: Gross Sales     | <cumulative>
            Row 1: Gross Expense   | <cumulative>
            Row 2: Gross Profit    | <cumulative>

        Each entry adds today's values on top of whatever is already stored.
        """
        data_path = os.path.join(OUTPUT_DIR, "data.docx")
        display_date = date.replace("_", "/")

        if not os.path.exists(data_path):
            # First-time creation — scan existing daily docs to seed cumulative totals
            # Today's entry is NOT included here; it gets added in the update block below
            cum_sales = 0
            cum_expense = 0
            cum_profit = 0
            for filename in os.listdir(OUTPUT_DIR):
                if filename.endswith(".docx") and filename != "data.docx":
                    try:
                        full_path = os.path.join(OUTPUT_DIR, filename)
                        doc = Document(full_path)
                        sales_val = int(doc.tables[1].rows[4].cells[1].text.strip())
                        net_text  = doc.tables[3].rows[0].cells[1].text.strip()
                        net_val   = int(net_text.partition("=")[2].strip())
                        expense_val = sales_val - net_val
                        cum_sales   += sales_val
                        cum_expense += expense_val
                        cum_profit  += net_val
                    except Exception as e:
                        print(f"[info] Skipping {filename}: {e}")

            netdoc = Document()
            netdoc.add_heading(f"Summary as on {display_date}", level=0)
            table = netdoc.add_table(rows=3, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Gross Sales"
            table.cell(0, 1).text = str(cum_sales)
            table.cell(1, 0).text = "Gross Expense"
            table.cell(1, 1).text = str(cum_expense)
            table.cell(2, 0).text = "Gross Profit"
            table.cell(2, 1).text = str(cum_profit)
            netdoc.save(data_path)

        # Load existing data.docx, update table values and heading date
        newdoc = Document(data_path)
        table2 = newdoc.tables[0]

        prev_sales   = int(table2.cell(0, 1).text.strip() or "0")
        prev_expense = int(table2.cell(1, 1).text.strip() or "0")
        prev_profit  = int(table2.cell(2, 1).text.strip() or "0")

        table2.cell(0, 1).text = str(prev_sales   + self.profit)
        table2.cell(1, 1).text = str(prev_expense + self.loss)
        table2.cell(2, 1).text = str(prev_profit  + net_day_profit)

        # Update heading: clear all runs via XML then write fresh text into one run
        for para in newdoc.paragraphs:
            if "Summary" in para.text:
                # Remove every <w:r> run element from the paragraph XML directly
                from docx.oxml.ns import qn as _qn
                for r in para._p.findall(_qn("w:r")):
                    para._p.remove(r)
                # Add a single clean run with the new date
                para.add_run(f"Summary as on {display_date}")
                break

        newdoc.save(data_path)

    def save(self, date):
        path = os.path.join(OUTPUT_DIR, f"{date}.docx")
        self.doc.save(path)